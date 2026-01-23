"""
Contract document generator
"""
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import settings


async def generate_contract_document(
    application_id: str,
    company_name: str,
    registration_number: str,
    country: str,
    contact_phone: str,
    selected_plan: str,
    num_etfs: int = 10
) -> Path:
    """
    Generate a Word document contract for the service agreement.
    Returns the path to the generated document.
    """
    # Ensure contracts directory exists
    contracts_dir = Path(settings.CONTRACTS_DIR) / "generated"
    contracts_dir.mkdir(parents=True, exist_ok=True)
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
    
    # Title
    title = doc.add_heading('ETF OWNERSHIP INTELLIGENCE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 212, 170)
    
    subtitle = doc.add_heading('Service Agreement', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Agreement date
    date_para = doc.add_paragraph()
    date_para.add_run(f'Date: {datetime.utcnow().strftime("%d %B %Y")}').bold = True
    
    doc.add_paragraph()
    
    # Parties
    doc.add_heading('1. PARTIES', level=2)
    
    doc.add_paragraph(
        'This Service Agreement ("Agreement") is entered into between:'
    )
    
    doc.add_paragraph(
        f'ETAnalytics Ltd.\n'
        f'Registered in Ireland\n'
        f'Registration Number: IE123456\n'
        f'Address: Dublin, Ireland\n'
        f'("Provider")'
    )
    
    doc.add_paragraph('AND')
    
    doc.add_paragraph(
        f'{company_name}\n'
        f'Registration Number: {registration_number}\n'
        f'Country: {country}\n'
        f'Phone: {contact_phone}\n'
        f'("Client")'
    )
    
    doc.add_paragraph()
    
    # Services
    doc.add_heading('2. SERVICES', level=2)
    
    doc.add_paragraph(
        'The Provider agrees to provide the Client with ETF ownership intelligence services, '
        'including but not limited to:'
    )
    
    services = [
        'Analysis of share register data to identify beneficial owners',
        'Entity matching and classification using proprietary database',
        'Disclosure request coordination (where applicable)',
        'Ownership reporting and analytics',
        'Access to the ET Analytics portal'
    ]
    
    for service in services:
        doc.add_paragraph(service, style='List Bullet')
    
    doc.add_paragraph()
    
    # Plan and Fees
    doc.add_heading('3. PLAN AND FEES', level=2)
    
    plan_details = {
        'starter': {
            'fee': '£80,000',
            'products': 'Up to 5 ETF products',
            'registers': '12 per year',
            'support': '24-hour support response'
        },
        'professional': {
            'fee': '£150,000',
            'products': 'Up to 15 ETF products',
            'registers': 'Unlimited',
            'support': '12-hour support response'
        },
        'enterprise': {
            'fee': 'Custom pricing',
            'products': 'Unlimited ETF products',
            'registers': 'Unlimited',
            'support': '4-hour support response'
        }
    }
    
    plan = plan_details.get(selected_plan.lower(), plan_details['professional'])
    
    doc.add_paragraph(f"Selected Plan: {selected_plan.upper()}")
    doc.add_paragraph(f"Annual Fee: {plan['fee']}")
    doc.add_paragraph(f"Coverage: {plan['products']}")
    doc.add_paragraph(f"Share Registers: {plan['registers']}")
    doc.add_paragraph(f"Support SLA: {plan['support']}")
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Fees are payable annually in advance. The Provider reserves the right to '
        'adjust fees upon renewal with 90 days notice.'
    )
    
    doc.add_paragraph()
    
    # Term and Termination
    doc.add_heading('4. TERM AND TERMINATION', level=2)
    
    doc.add_paragraph(
        'This Agreement commences on the date of signing and continues for an initial '
        'term of 12 months. The Agreement will automatically renew for successive 12-month '
        'periods unless either party provides written notice of non-renewal at least 90 days '
        'prior to the end of the current term.'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Either party may terminate this Agreement for material breach upon 30 days written '
        'notice if the breach remains uncured.'
    )
    
    doc.add_paragraph()
    
    # Data and Confidentiality
    doc.add_heading('5. DATA AND CONFIDENTIALITY', level=2)
    
    doc.add_paragraph(
        'The Provider agrees to:'
    )
    
    data_terms = [
        'Process share register data solely for the purposes of this Agreement',
        'Maintain appropriate technical and organizational security measures',
        'Comply with applicable data protection laws including GDPR',
        'Not disclose Client data to third parties without prior written consent',
        'Return or destroy Client data upon termination'
    ]
    
    for term in data_terms:
        doc.add_paragraph(term, style='List Bullet')
    
    doc.add_paragraph()
    
    # Limitation of Liability
    doc.add_heading('6. LIMITATION OF LIABILITY', level=2)
    
    doc.add_paragraph(
        'To the maximum extent permitted by law, the Provider\'s total liability under this '
        'Agreement shall not exceed the fees paid by the Client in the 12 months preceding '
        'the claim. The Provider shall not be liable for any indirect, incidental, or '
        'consequential damages.'
    )
    
    doc.add_paragraph()
    
    # Signatures
    doc.add_heading('7. SIGNATURES', level=2)
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 40)
    doc.add_paragraph('ETAnalytics Ltd.')
    doc.add_paragraph('Name:')
    doc.add_paragraph('Title:')
    doc.add_paragraph('Date:')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph('_' * 40)
    doc.add_paragraph(f'{company_name}')
    doc.add_paragraph('Name:')
    doc.add_paragraph('Title:')
    doc.add_paragraph('Date:')
    
    # Save document
    safe_name = company_name.replace(' ', '_').replace('/', '-')[:50]
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"ETAnalytics_Service_Agreement_{safe_name}_{timestamp}.docx"
    filepath = contracts_dir / filename
    
    doc.save(filepath)
    
    return filepath
